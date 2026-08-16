import io
import re
import os
import shutil
import logging
import numpy as np
from dataclasses import dataclass, field
from typing import List, Dict, Any, Optional

logger = logging.getLogger(__name__)

# Standard resume section patterns
SECTION_PATTERNS = {
    "SUMMARY": r"\b(summary|objective|profile|about\s+me)\b",
    "EXPERIENCE": r"\b(experience|work\s+experience|employment|work\s+history|internships?)\b",
    "PROJECTS": r"\b(projects|key\s+projects|personal\s+projects|portfolio)\b",
    "EDUCATION": r"\b(education|academic\s+background|qualifications|degrees?)\b",
    "SKILLS": r"\b(skills|technical\s+skills|core\s+competencies|technologies|tools)\b",
    "CERTIFICATIONS": r"\b(certifications?|licenses?|credentials?)\b"
}

def detect_tesseract_binary() -> Optional[str]:
    """Finds Tesseract OCR executable on Windows or system PATH if available."""
    path = shutil.which("tesseract")
    if path:
        return path
    win_paths = [
        r"C:\Program Files\Tesseract-OCR\tesseract.exe",
        r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
        os.path.expanduser(r"~\AppData\Local\Programs\Tesseract-OCR\tesseract.exe")
    ]
    for wp in win_paths:
        if os.path.exists(wp):
            return wp
    return None

def detect_resume_sections(text: str) -> List[str]:
    """Detects standard resume section headings in normalized text."""
    if not text:
        return []
    text_lower = text.lower()
    found = []
    for sec_name, pattern in SECTION_PATTERNS.items():
        if re.search(pattern, text_lower):
            found.append(sec_name)
    return found

@dataclass
class ExtractionQualityResult:
    score: float  # 0.0 to 100.0
    status: str   # "PASS", "FAIL", "SCANNED_PDF", "ENCRYPTED_PDF", "INVALID_PDF", "EMPTY_PDF"
    warnings: List[str] = field(default_factory=list)
    char_count: int = 0
    word_count: int = 0
    alpha_ratio: float = 0.0
    printable_ratio: float = 0.0
    symbol_ratio: float = 0.0
    has_email: bool = False
    has_url: bool = False
    has_date: bool = False
    detected_sections: List[str] = field(default_factory=list)
    reasons: List[str] = field(default_factory=list)

@dataclass
class ResumeExtractionResult:
    filename: str
    file_type: str
    raw_text: str
    normalized_text: str
    character_count: int
    page_count: int
    extraction_method: str  # "pymupdf", "pdfplumber", "pypdf", "ocr", "python-docx", "utf-8"
    quality: ExtractionQualityResult
    page_diagnostics: List[Dict[str, Any]] = field(default_factory=list)

def normalize_text(text: str) -> str:
    """
    Normalizes extracted resume text while strictly preserving punctuation critical for tech stacks
    (e.g., C++, C#, .NET, Node.js, React.js, FastAPI, PostgreSQL, AWS, GCP, Azure, PyTorch).
    """
    if not text:
        return ""

    # Remove null bytes and control characters except newlines/tabs
    text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', text)

    # Normalize unicode whitespace to standard spaces
    text = re.sub(r'[\u2000-\u200B\u00A0]', ' ', text)

    # Normalize excessive blank lines while preserving section breaks
    lines = [line.strip() for line in text.splitlines()]
    normalized_lines = []
    prev_blank = False

    for line in lines:
        if not line:
            if not prev_blank:
                normalized_lines.append("")
                prev_blank = True
        else:
            normalized_lines.append(line)
            prev_blank = False

    return "\n".join(normalized_lines).strip()

def validate_extracted_text(text: str, page_count: int = 1) -> ExtractionQualityResult:
    """
    Evaluates text quality using multi-signal metrics: character count, word count,
    printable ASCII ratio, alphabetic ratio, symbol noise ratio, email/date detection,
    and resume section detection.
    """
    warnings = []
    reasons = []
    normalized = normalize_text(text)
    total_chars = len(normalized)

    has_email = bool(re.search(r'[\w\.-]+@[\w\.-]+\.\w+', normalized))
    has_url = bool(re.search(r'(github\.com|linkedin\.com|http|https|\b\w+\.(io|com|org|dev)\b)', normalized, re.IGNORECASE))
    has_date = bool(re.search(r'\b(20\d\d|19\d\d|present|jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)\b', normalized, re.IGNORECASE))
    sections = detect_resume_sections(normalized)

    words = [w for w in re.split(r'\s+', normalized) if len(w) > 1]
    word_count = len(words)

    if total_chars < 50 or word_count < 10:
        warnings.append("Extracted text is empty or too short (<50 chars). File may be image-based or scanned.")
        reasons.append("insufficient_readable_text")
        return ExtractionQualityResult(
            score=0.0,
            status="SCANNED_PDF" if page_count > 0 else "FAIL",
            warnings=warnings,
            char_count=total_chars,
            word_count=word_count,
            alpha_ratio=0.0,
            printable_ratio=0.0,
            symbol_ratio=0.0,
            has_email=has_email,
            has_url=has_url,
            has_date=has_date,
            detected_sections=sections,
            reasons=reasons
        )

    # Calculate character ratios
    alpha_chars = sum(c.isalpha() for c in normalized)
    digit_chars = sum(c.isdigit() for c in normalized)
    space_chars = sum(c.isspace() for c in normalized)
    printable_chars = sum(32 <= ord(c) <= 126 or c in '\n\r\t' for c in normalized)
    symbol_chars = total_chars - (alpha_chars + digit_chars + space_chars)

    alpha_ratio = alpha_chars / total_chars
    printable_ratio = printable_chars / total_chars
    symbol_ratio = symbol_chars / total_chars

    score = 80.0  # Base score for non-empty text

    # Add signal bonuses
    if has_email:
        score += 5.0
    if has_date:
        score += 5.0
    if len(sections) >= 2:
        score += 10.0

    # Penalties for noise
    if printable_ratio < 0.85:
        score -= (1.0 - printable_ratio) * 100
        warnings.append(f"Low printable ASCII character ratio ({printable_ratio:.1%}).")
        reasons.append("low_printable_ratio")

    if alpha_ratio < 0.35:
        score -= (0.35 - alpha_ratio) * 120
        warnings.append(f"Low alphabetic character ratio ({alpha_ratio:.1%}). Text contains excessive noise.")
        reasons.append("low_alphabetic_ratio")

    if symbol_ratio > 0.25:
        score -= symbol_ratio * 100
        warnings.append(f"High random symbol density ({symbol_ratio:.1%}).")
        reasons.append("high_symbol_noise")

    garbage_tokens = re.findall(r'\b[A-Z]{1,3}[%*#&$][A-Z0-9%*#&$]{1,4}\b', normalized)
    if len(garbage_tokens) > 2:
        score -= min(30, len(garbage_tokens) * 5)
        warnings.append(f"Detected {len(garbage_tokens)} corrupted stream tokens.")
        reasons.append("corrupted_font_stream_tokens")

    score = max(0.0, min(100.0, round(score, 1)))

    status = "PASS"
    if score < 50.0 or alpha_ratio < 0.30:
        status = "FAIL"

    return ExtractionQualityResult(
        score=score,
        status=status,
        warnings=warnings,
        char_count=total_chars,
        word_count=word_count,
        alpha_ratio=round(alpha_ratio, 3),
        printable_ratio=round(printable_ratio, 3),
        symbol_ratio=round(symbol_ratio, 3),
        has_email=has_email,
        has_url=has_url,
        has_date=has_date,
        detected_sections=sections,
        reasons=reasons
    )

def extract_pdf_text(contents: bytes, filename: str) -> ResumeExtractionResult:
    """
    Multi-strategy PDF text extraction:
    - Signature Header Check (%PDF-)
    - Encryption Detection
    - Strategy 1: PyMuPDF (fitz/pymupdf) with text & blocks extraction
    - Strategy 2: pdfplumber layout parser
    - Strategy 3: pypdf layout parser
    - Strategy 4: Python-native RapidOCR + Tesseract OCR fallback for scanned/image PDFs
    """
    # 1. Signature Header Check
    if not contents or len(contents) == 0:
        q = ExtractionQualityResult(
            score=0.0,
            status="EMPTY_PDF",
            warnings=["Uploaded file is empty (0 bytes)."],
            reasons=["empty_file"]
        )
        return ResumeExtractionResult(
            filename=filename,
            file_type="application/pdf",
            raw_text="",
            normalized_text="",
            character_count=0,
            page_count=0,
            extraction_method="none",
            quality=q
        )

    if b"%PDF-" not in contents[:1024]:
        q = ExtractionQualityResult(
            score=0.0,
            status="INVALID_PDF",
            warnings=["Uploaded file does not have a valid PDF header signature (%PDF-). The file may be corrupt."],
            reasons=["invalid_pdf_header"]
        )
        return ResumeExtractionResult(
            filename=filename,
            file_type="application/pdf",
            raw_text="",
            normalized_text="",
            character_count=0,
            page_count=0,
            extraction_method="none",
            quality=q
        )

    raw_text = ""
    page_count = 0
    method = "pymupdf"
    page_diagnostics = []
    is_encrypted = False

    # 2. Encryption Check
    try:
        import pymupdf
        doc = pymupdf.open(stream=contents, filetype="pdf")
        page_count = len(doc)
        if doc.is_encrypted or doc.needs_pass:
            is_encrypted = True
    except Exception:
        pass

    if is_encrypted:
        q = ExtractionQualityResult(
            score=0.0,
            status="ENCRYPTED_PDF",
            warnings=["This PDF is password-protected or encrypted. Please upload an unlocked copy."],
            reasons=["encrypted_pdf"]
        )
        return ResumeExtractionResult(
            filename=filename,
            file_type="application/pdf",
            raw_text="",
            normalized_text="",
            character_count=0,
            page_count=page_count,
            extraction_method="none",
            quality=q
        )

    # 3. Strategy 1: PyMuPDF (fitz)
    try:
        import pymupdf
        doc = pymupdf.open(stream=contents, filetype="pdf")
        page_count = len(doc)
        pages_text = []
        for i, page in enumerate(doc):
            txt = page.get_text("text") or ""
            if not txt.strip():
                blocks = page.get_text("blocks")
                txt = "\n".join([b[4] for b in blocks if len(b) >= 5 and isinstance(b[4], str)])
            pages_text.append(txt)
            page_diagnostics.append({
                "page": i + 1,
                "pymupdf_chars": len(txt.strip()),
                "has_images": len(page.get_images()) > 0
            })
        raw_text = "\n\n".join(pages_text)
    except Exception as e:
        logger.warning(f"PyMuPDF extraction failed for {filename}: {e}")

    quality = validate_extracted_text(raw_text, page_count=page_count)

    # 4. Strategy 2: Fallback to pdfplumber
    if quality.status != "PASS":
        try:
            import pdfplumber
            with pdfplumber.open(io.BytesIO(contents)) as pdf:
                page_count = len(pdf.pages)
                plumber_pages = []
                for i, page in enumerate(pdf.pages):
                    txt = page.extract_text(layout=True) or page.extract_text() or ""
                    plumber_pages.append(txt)
                    if i < len(page_diagnostics):
                        page_diagnostics[i]["pdfplumber_chars"] = len(txt.strip())
                plumber_text = "\n\n".join(plumber_pages)
                plumber_quality = validate_extracted_text(plumber_text, page_count=page_count)
                if plumber_quality.score > quality.score:
                    raw_text = plumber_text
                    quality = plumber_quality
                    method = "pdfplumber"
        except Exception as e:
            logger.warning(f"pdfplumber extraction failed for {filename}: {e}")

    # 5. Strategy 3: Fallback to pypdf
    if quality.status != "PASS":
        try:
            import pypdf
            reader = pypdf.PdfReader(io.BytesIO(contents))
            pypdf_pages = []
            for i, p in enumerate(reader.pages):
                txt = p.extract_text(extraction_mode="layout") or p.extract_text() or ""
                pypdf_pages.append(txt)
                if i < len(page_diagnostics):
                    page_diagnostics[i]["pypdf_chars"] = len(txt.strip())
            pypdf_text = "\n\n".join(pypdf_pages)
            pypdf_quality = validate_extracted_text(pypdf_text, page_count=page_count)
            if pypdf_quality.score > quality.score:
                raw_text = pypdf_text
                quality = pypdf_quality
                method = "pypdf"
        except Exception as e:
            logger.warning(f"pypdf extraction failed for {filename}: {e}")

    # 6. Strategy 4: Pure Python RapidOCR + Tesseract Fallback for Scanned/Image PDFs
    if quality.status != "PASS":
        try:
            import pymupdf
            from rapidocr_onnxruntime import RapidOCR
            ocr_engine = RapidOCR()
            doc = pymupdf.open(stream=contents, filetype="pdf")
            ocr_pages = []

            for i, page in enumerate(doc):
                pix = page.get_pixmap(dpi=300)
                img_np = np.frombuffer(pix.samples, dtype=np.uint8).reshape(pix.h, pix.w, pix.n)
                result, elapse = ocr_engine(img_np)
                if result:
                    page_txt = "\n".join([res[1] for res in result if res and len(res) >= 2 and res[1]])
                    ocr_pages.append(page_txt)

            ocr_text = "\n\n".join(ocr_pages)
            ocr_quality = validate_extracted_text(ocr_text, page_count=page_count)

            if ocr_quality.status == "PASS" or ocr_quality.score > quality.score:
                raw_text = ocr_text
                quality = ocr_quality
                method = "ocr"
                logger.info(f"RapidOCR successful for {filename}: extracted {len(ocr_text)} chars.")
        except Exception as rapid_err:
            logger.warning(f"RapidOCR fallback error for {filename}: {rapid_err}")
            # Try Tesseract if RapidOCR was unavailable
            tess_cmd = detect_tesseract_binary()
            if tess_cmd:
                try:
                    import pytesseract
                    import pymupdf
                    from PIL import Image

                    pytesseract.pytesseract.tesseract_cmd = tess_cmd
                    doc = pymupdf.open(stream=contents, filetype="pdf")
                    ocr_pages = []

                    for page in doc:
                        pix = page.get_pixmap(dpi=300)
                        img = Image.open(io.BytesIO(pix.tobytes("png")))
                        ocr_txt = pytesseract.image_to_string(img)
                        ocr_pages.append(ocr_txt)

                    ocr_text = "\n\n".join(ocr_pages)
                    ocr_quality = validate_extracted_text(ocr_text, page_count=page_count)

                    if ocr_quality.status == "PASS" or ocr_quality.score > quality.score:
                        raw_text = ocr_text
                        quality = ocr_quality
                        method = "ocr"
                except Exception as ocr_err:
                    logger.error(f"Tesseract OCR error for {filename}: {ocr_err}")

    normalized = normalize_text(raw_text)

    # Detailed per-page diagnostics logging
    for diag in page_diagnostics:
        logger.info(f"PDF DIAGNOSTIC | FILE: {filename} | PAGE {diag['page']}: PyMuPDF chars={diag.get('pymupdf_chars', 0)} | pdfplumber chars={diag.get('pdfplumber_chars', 0)} | pypdf chars={diag.get('pypdf_chars', 0)} | has_images={diag.get('has_images', False)}")

    logger.info(f"PDF SUMMARY | FILE: {filename} | METHOD: {method} | PAGES: {page_count} | TOTAL CHARS: {len(normalized)} | SCORE: {quality.score}% | STATUS: {quality.status}")

    return ResumeExtractionResult(
        filename=filename,
        file_type="application/pdf",
        raw_text=raw_text,
        normalized_text=normalized,
        character_count=len(normalized),
        page_count=page_count,
        extraction_method=method,
        quality=quality,
        page_diagnostics=page_diagnostics
    )

def extract_docx_text(contents: bytes, filename: str) -> ResumeExtractionResult:
    """
    Extracts text from DOCX files preserving headings, paragraphs, bullet points, and tables.
    """
    raw_text = ""
    method = "python-docx"

    try:
        import docx
        doc = docx.Document(io.BytesIO(contents))
        parts = []

        # Extract paragraphs & headings
        for p in doc.paragraphs:
            if p.text and p.text.strip():
                parts.append(p.text.strip())

        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_cells:
                    parts.append(" | ".join(row_cells))

        raw_text = "\n".join(parts)
    except Exception as e:
        logger.error(f"DOCX extraction failed for {filename}: {e}")
        raw_text = ""

    normalized = normalize_text(raw_text)
    quality = validate_extracted_text(normalized, page_count=1)
    logger.info(f"FILE: {filename} | TYPE: DOCX | METHOD: {method} | CHARS: {len(normalized)} | SCORE: {quality.score}% | STATUS: {quality.status}")

    return ResumeExtractionResult(
        filename=filename,
        file_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        raw_text=raw_text,
        normalized_text=normalized,
        character_count=len(normalized),
        page_count=1,
        extraction_method=method,
        quality=quality,
        page_diagnostics=[{"page": 1, "chars": len(normalized)}]
    )

def extract_txt_text(contents: bytes, filename: str) -> ResumeExtractionResult:
    """
    Safely decodes TXT files using UTF-8 with fallback encodings.
    """
    raw_text = ""
    method = "utf-8"

    for enc in ["utf-8", "latin-1", "cp1252"]:
        try:
            raw_text = contents.decode(enc)
            method = enc
            break
        except UnicodeDecodeError:
            continue

    normalized = normalize_text(raw_text)
    quality = validate_extracted_text(normalized, page_count=1)
    logger.info(f"FILE: {filename} | TYPE: TXT | METHOD: {method} | CHARS: {len(normalized)} | SCORE: {quality.score}% | STATUS: {quality.status}")

    return ResumeExtractionResult(
        filename=filename,
        file_type="text/plain",
        raw_text=raw_text,
        normalized_text=normalized,
        character_count=len(normalized),
        page_count=1,
        extraction_method=method,
        quality=quality,
        page_diagnostics=[{"page": 1, "chars": len(normalized)}]
    )

def extract_and_validate_resume(contents: bytes, filename: str) -> ResumeExtractionResult:
    """
    Main entry point for extracting and validating resume file text.
    """
    fn_lower = filename.lower()
    if fn_lower.endswith(".pdf"):
        return extract_pdf_text(contents, filename)
    elif fn_lower.endswith(".docx"):
        return extract_docx_text(contents, filename)
    elif fn_lower.endswith(".txt"):
        return extract_txt_text(contents, filename)
    else:
        q = ExtractionQualityResult(
            score=0.0, 
            status="FAIL", 
            warnings=["Unsupported file extension. Only PDF, DOCX, and TXT are supported."]
        )
        return ResumeExtractionResult(
            filename=filename,
            file_type="unknown",
            raw_text="",
            normalized_text="",
            character_count=0,
            page_count=0,
            extraction_method="none",
            quality=q
        )

def validate_llm_profile_output(profile_dict: Dict[str, Any], source_text: str) -> tuple[Dict[str, Any], float]:
    """
    Post-LLM anti-hallucination validator. Cross-references extracted fields against source resume text.
    Strips unevidenced skills, hallucinated placeholder emails (like candidate@skillproof.io), and computes
    a trustworthy parsing confidence score (0-100%).
    """
    source_lower = source_text.lower()
    verified_profile = dict(profile_dict)
    validation_flags = []
    confidence_points = 100.0

    # 1. Validate email
    email = verified_profile.get("email")
    if email:
        email_clean = email.strip().lower()
        if "candidate@skillproof" in email_clean or "jane.doe" in email_clean or "john.doe" in email_clean:
            verified_profile["email"] = None
            validation_flags.append("Removed placeholder email.")
            confidence_points -= 10.0
        elif email_clean not in source_lower:
            verified_profile["email"] = None
            validation_flags.append(f"Removed unevidenced email '{email}'.")
            confidence_points -= 15.0

    # 2. Validate candidate name
    name = verified_profile.get("name")
    if name:
        name_clean = name.strip()
        if name_clean.lower() in ["jane doe", "john doe", "z d%", "candidate profile", "resume"]:
            verified_profile["name"] = "Candidate Profile"
            confidence_points -= 15.0

    # 3. Validate skills against source text
    raw_skills = verified_profile.get("skills") or []
    verified_skills = []
    for skill in raw_skills:
        skill_clean = skill.strip()
        if not skill_clean:
            continue
        skill_lower = skill_clean.lower()
        regex_pattern = r'\b' + re.escape(skill_lower) + r'\b'
        if skill_lower in source_lower or re.search(regex_pattern, source_lower):
            verified_skills.append(skill_clean)
        elif skill_lower == "go" and not re.search(r'\b(golang|go)\b', source_lower):
            validation_flags.append("Removed hallucinated skill 'Go'.")
            confidence_points -= 10.0
        elif len(skill_clean) > 2 and skill_lower in source_lower:
            verified_skills.append(skill_clean)
        else:
            if len(skill_clean) >= 4 and any(part in source_lower for part in skill_lower.split()):
                verified_skills.append(skill_clean)
            else:
                validation_flags.append(f"Removed unevidenced skill '{skill_clean}'.")
                confidence_points -= 5.0

    verified_profile["skills"] = verified_skills

    confidence_score = max(0.0, min(100.0, round(confidence_points, 1)))
    verified_profile["parse_confidence_score"] = confidence_score
    verified_profile["validation_flags"] = validation_flags

    return verified_profile, confidence_score


def parse_resume_text_to_profile(text: str) -> Dict[str, Any]:
    """
    Direct section-based resume parser extracting Projects, Education, Experience, Skills,
    and Personal Info strictly from the supplied resume text without mock placeholders.
    """
    if not text:
        return {
            "name": "Candidate Profile",
            "email": None,
            "summary": None,
            "skills": [],
            "experience": [],
            "projects": [],
            "education": []
        }

    lines = [l.strip() for l in text.splitlines() if l.strip()]

    # 1. Candidate Name
    candidate_name = "Candidate Profile"
    for line in lines[:5]:
        clean = re.sub(r'^(name|resume|cv):', '', line, flags=re.IGNORECASE).strip()
        if 2 <= len(clean.split()) <= 4 and not any(c.isdigit() for c in clean) and "@" not in clean and "http" not in clean and "/" not in clean:
            candidate_name = clean.title()
            break

    # 2. Email
    email_match = re.search(r'[\w\.-]+@[\w\.-]+\.\w+', text)
    extracted_email = email_match.group(0) if email_match else None

    # 3. GitHub Username & URL
    github_match = re.search(r'github\.com/([\w-]+)', text, re.IGNORECASE)
    extracted_github_username = github_match.group(1) if github_match else None
    github_url = f"https://github.com/{extracted_github_username}" if extracted_github_username else None

    # 4. LinkedIn URL
    linkedin_match = re.search(r'linkedin\.com/in/([\w-]+)', text, re.IGNORECASE)
    extracted_linkedin_url = f"https://linkedin.com/in/{linkedin_match.group(1)}" if linkedin_match else None

    # 4. Partition lines into section blocks
    section_headers = {
        "SUMMARY": r"^(summary|objective|profile|about\s+me)$",
        "EDUCATION": r"^(education|academic\s+background|qualifications|academic\s+credentials)$",
        "SKILLS": r"^(skills|technical\s+skills|core\s+competencies|technologies)$",
        "EXPERIENCE": r"^(experience|work\s+experience|employment|work\s+history|internships)$",
        "PROJECTS": r"^(projects|key\s+projects|personal\s+projects|portfolio)$",
        "CERTIFICATIONS": r"^(certifications?|achievements?|certifications\s+&\s+achievements|honors)$"
    }

    sections: Dict[str, List[str]] = {
        "HEADER": [],
        "SUMMARY": [],
        "EDUCATION": [],
        "SKILLS": [],
        "EXPERIENCE": [],
        "PROJECTS": [],
        "CERTIFICATIONS": []
    }

    current_sec = "HEADER"

    for line in text.splitlines():
        l_clean = line.strip()
        if not l_clean:
            continue

        matched_sec = None
        for sec_name, pattern in section_headers.items():
            if re.match(pattern, l_clean, re.IGNORECASE):
                matched_sec = sec_name
                break

        if matched_sec:
            current_sec = matched_sec
        else:
            sections[current_sec].append(l_clean)

    # -------------------------------------------------------------
    # Parse EDUCATION directly from text
    # -------------------------------------------------------------
    education_list = []
    edu_lines = sections["EDUCATION"]
    if edu_lines:
        curr_edu = {"degree": "", "institution": "", "year": "", "gpa": "", "details": []}
        for el in edu_lines:
            el_clean = el.lstrip("•-* ").strip()
            el_lower = el_clean.lower()

            if any(d in el_lower for d in ["b.tech", "b.e.", "b.s.", "m.s.", "bachelor", "master", "degree", "phd", "diploma"]):
                if curr_edu["degree"]:
                    education_list.append(curr_edu)
                    curr_edu = {"degree": "", "institution": "", "year": "", "gpa": "", "details": []}
                curr_edu["degree"] = el_clean
            elif any(dt in el_lower for dt in ["2023", "2024", "2025", "2026", "2027", "2022", "2021", "2020", "2019", "present", "jan", "feb", "mar", "apr", "may", "jun", "jul", "aug", "sep", "oct", "nov", "dec"]) and any(c.isdigit() for c in el_clean) and len(el_clean) < 35:
                curr_edu["year"] = el_clean
            elif any(inst in el_lower for inst in ["university", "institute", "college", "school", "academy"]):
                curr_edu["institution"] = el_clean
            elif "cgpa" in el_lower or "gpa" in el_lower or "backlog" in el_lower:
                curr_edu["gpa"] = el_clean
            elif "coursework" in el_lower or el.startswith("•") or el.startswith("-"):
                curr_edu["details"].append(el_clean)
            elif not curr_edu["institution"] and not curr_edu["degree"]:
                curr_edu["institution"] = el_clean

        if curr_edu["degree"] or curr_edu["institution"] or curr_edu["gpa"]:
            education_list.append(curr_edu)

    # -------------------------------------------------------------
    # Parse PROJECTS directly from text
    # -------------------------------------------------------------
    projects_list = []
    proj_lines = sections["PROJECTS"]

    if proj_lines:
        curr_proj = None
        for pl in proj_lines:
            pl_clean = pl.lstrip("•-* ").strip()
            if not pl_clean:
                continue

            pl_lower = pl_clean.lower()
            if pl_lower in ["github", "github ->", "github →", "github link"]:
                continue

            is_explicit_bullet = pl.startswith("•") or pl.startswith("-") or pl.startswith("*")
            is_title = ("–" in pl_clean or "—" in pl_clean or (len(pl_clean) < 60 and not pl_clean.endswith("."))) and not is_explicit_bullet

            if is_title and (curr_proj is None or curr_proj.get("description")):
                if curr_proj and curr_proj.get("name"):
                    projects_list.append(curr_proj)
                curr_proj = {
                    "name": pl_clean,
                    "description": "",
                    "technologies": [],
                    "repo_url": github_url
                }
            elif curr_proj and not curr_proj["description"] and not is_explicit_bullet and ("," in pl_clean or len(pl_clean.split()) <= 6):
                techs = [t.strip() for t in pl_clean.split(",") if t.strip()]
                curr_proj["technologies"] = techs
            elif curr_proj:
                if curr_proj["description"]:
                    curr_proj["description"] += " " + pl_clean
                else:
                    curr_proj["description"] = pl_clean

        if curr_proj and curr_proj.get("name"):
            projects_list.append(curr_proj)

    # -------------------------------------------------------------
    # Parse EXPERIENCE directly from text
    # -------------------------------------------------------------
    experience_list = []
    exp_lines = sections["EXPERIENCE"]

    if exp_lines:
        curr_exp = None
        for xl in exp_lines:
            xl_clean = xl.lstrip("•-* ").strip()
            if not xl_clean:
                continue

            xl_lower = xl_clean.lower()
            is_explicit_bullet = xl.startswith("•") or xl.startswith("-") or xl.startswith("*")
            is_date = any(dt in xl_lower for dt in ["2023", "2024", "2025", "2026", "2027", "present", "jun", "jul", "aug", "sep", "oct", "nov", "dec", "jan", "feb", "mar", "apr", "may"]) and any(c.isdigit() for c in xl_clean) and len(xl_clean) < 35
            is_title = not is_explicit_bullet and xl_clean[0].isupper() and len(xl_clean) < 65 and not xl_clean.endswith(".") and not any(kw in xl_lower for kw in ["accuracy", "execution", "delivered", "pipeline", "workshops", "bottleneck"])

            if is_date:
                if curr_exp:
                    curr_exp["dates"] = xl_clean
            elif is_title and (curr_exp is None or len(curr_exp.get("description", [])) > 0):
                if curr_exp and curr_exp.get("title"):
                    experience_list.append(curr_exp)
                curr_exp = {
                    "title": xl_clean,
                    "company": "Organization",
                    "dates": "Recent",
                    "description": [],
                    "technologies": []
                }
            elif curr_exp and not curr_exp["description"] and not is_explicit_bullet:
                if curr_exp["company"] == "Organization":
                    curr_exp["company"] = xl_clean
                else:
                    curr_exp["company"] += " - " + xl_clean
            elif curr_exp:
                if is_explicit_bullet or not curr_exp["description"]:
                    curr_exp["description"].append(xl_clean)
                else:
                    curr_exp["description"][-1] += " " + xl_clean

        if curr_exp and curr_exp.get("title"):
            experience_list.append(curr_exp)

    # -------------------------------------------------------------
    # Parse SKILLS directly from text
    # -------------------------------------------------------------
    skills_list = []
    skill_lines = sections["SKILLS"]
    if skill_lines:
        for sl in skill_lines:
            sl_clean = sl.lstrip("•-* ").strip()
            if ":" in sl_clean:
                sl_clean = sl_clean.split(":", 1)[1].strip()
            for part in sl_clean.split(","):
                part_clean = part.strip()
                if part_clean and part_clean not in skills_list and len(part_clean) < 30:
                    skills_list.append(part_clean)

    if not skills_list:
        common_skills = [
            "Python", "JavaScript", "TypeScript", "Java", "C++", "C#", "Go", "Rust",
            "FastAPI", "Django", "Flask", "React", "React Native", "Next.js", "Node.js", "Express",
            "Docker", "PostgreSQL", "MySQL", "MongoDB", "Redis", "Git", "REST", "HTML5", "CSS3"
        ]
        skills_list = [s for s in common_skills if re.search(r'\b' + re.escape(s) + r'\b', text, re.IGNORECASE)]

    summary_text = " ".join(sections["SUMMARY"]) if sections["SUMMARY"] else text[:300]

    return {
        "name": candidate_name,
        "email": extracted_email,
        "github_username": extracted_github_username,
        "linkedin_url": extracted_linkedin_url,
        "summary": summary_text,
        "skills": skills_list,
        "experience": experience_list,
        "projects": projects_list,
        "education": education_list
    }
